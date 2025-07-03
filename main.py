# ✅ Updated Streamlit + LangGraph Agentic AI App

import streamlit as st
from langgraph.graph import StateGraph, END
from typing import TypedDict, Optional, List
from langchain_core.runnables import Runnable
#from langchain_community.llms import HuggingFaceEndpoint
from huggingface_hub import InferenceClient
from serpapi import GoogleSearch
import vanna as vn
from vanna.remote import VannaDefault
from docx import Document
import tempfile
import os
from dotenv import load_dotenv
import json
import re
from openai import OpenAI
import openai
import pandas as pd
import requests
import sqlite3

load_dotenv()

# ---- Vanna Setup ----
vanna_api_key = os.getenv("vanna_api_key")
vanna_model_name = os.getenv("vanna_model_name")
vn_model = VannaDefault(model=vanna_model_name, api_key=vanna_api_key)
vn_model.connect_to_sqlite('/Users/vikalp.bhatnagar/Downloads/Actuarial_PC.db')

# ---- Config ----
#openai.api_key = os.getenv("TOGETHER_API_KEY")
#openai.base_url = "https://api.together.xyz"
serpapi_key = os.getenv("SERPAPI_API_KEY")

# ---Open AI LLM---

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def call_llm(prompt: str) -> str:
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an intelligent AI assistant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=500
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"OpenAI call failed: {e}"

# # ---- Together AI LLM ----

# def call_llm(prompt: str) -> str:
#     together_api_key = os.getenv("TOGETHER_API_KEY")
#     url = "https://api.together.xyz/v1/chat/completions"
#     headers = {
#         "Authorization": f"Bearer {together_api_key}",
#         "Content-Type": "application/json"
#     }
#     data = {
#         "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
#         "messages": [
#             {"role": "system", "content": "You are an intelligent AI assistant."},
#             {"role": "user", "content": prompt}
#         ],
#         "temperature": 0.1,
#         "max_tokens": 500
#     }

#     response = requests.post(url, headers=headers, json=data)
#     response.raise_for_status()
#     return response.json()['choices'][0]['message']['content'].strip()

# ---- Define LangGraph State ----
class GraphState(TypedDict):
    user_prompt: str
    doc_loaded: bool
    document_path: Optional[str]
    vanna_prompt: Optional[str]
    fuzzy_prompt: Optional[str]
    route: Optional[str]
    sql_result: Optional[pd.DataFrame]
    web_links: Optional[List[str]]
    updated_doc_path: Optional[str]


def get_schema_description(db_path: str) -> str:
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    schema_str = ""
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
    tables = cursor.fetchall()

    for table_name, in tables:
        cursor.execute(f"PRAGMA table_info({table_name});")
        cols = cursor.fetchall()
        col_names = [col[1] for col in cols]
        schema_str += f"\n- {table_name}: columns = {', '.join(col_names)}"

    conn.close()
    return schema_str.strip()


# ---- Router Node (with prompt generation) ----
class RouterNode(Runnable):
    def invoke(self, state: GraphState, config=None) -> GraphState:
        doc_flag = "yes" if state['doc_loaded'] else "no"

        schema = get_schema_description('/Users/vikalp.bhatnagar/Downloads/Actuarial_PC.db')

        router_prompt = f"""
        You are an intelligent routing agent. Your job is to:
        1. Choose one of the paths: "sql", "search", or "document" based on the user prompt.
        2. Choose:
        - "sql" if the user is asking a question about structured insurance data or something that can be answered from the following database schema:
            {schema}

        - "document" ONLY if a document is uploaded (Document Uploaded = yes) AND the question involves updating/reading a document.

        - "search" if it's a general query, involves latest events, external info, or cannot be answered from structured data.

        3. If the route is "sql" or "document", include:
        - "vanna_prompt": an SQL-style question to query structured data.

        4. If the route is "document", also include:
        - "fuzzy_prompt": a natural language description of the header or table to update.

        5. If the route is "search", DO NOT include vanna_prompt or fuzzy_prompt.

        Return output strictly in valid JSON format.

        Examples:

        For SQL:
        {{
            "route": "sql",
            "vanna_prompt": "SELECT * FROM claims WHERE loss_type = 'war'"
        }}

        For Document:
        {{
            "route": "document",
            "vanna_prompt": "SELECT policy_id, total_loss FROM policies WHERE year = 2024",
            "fuzzy_prompt": "Update the table under 'Loss Overview' for 2024"
        }}

        For Search:
        {{
            "route": "search"
        }}

        User Prompt: "{state['user_prompt']}"
        Document Uploaded: {doc_flag}
        """

        try:
            response = call_llm(router_prompt)
            st.write("[RouterNode] Raw LLM response:", response)  # 🔍 Show for debugging

            match = re.search(r'{.*}', response, re.DOTALL)
            if match:
                parsed = json.loads(match.group())
            else:
                st.warning("LLM did not return valid JSON. Routing to 'search' as fallback.")
                parsed = {"route": "search"}

        except Exception as e:
            st.error(f"[RouterNode] LLM call failed: {e}")
            parsed = {"route": "search"}

        # ✅ Additional safety check — avoid "document" route if no doc uploaded
        if parsed.get("route") == "document" and not state['doc_loaded']:
            parsed["route"] = "search"
            parsed["vanna_prompt"] = None
            parsed["fuzzy_prompt"] = None
            st.warning("Document route selected but no document uploaded. Falling back to search.")

        return {
            **state,
            "route": parsed.get("route"),
            "vanna_prompt": parsed.get("vanna_prompt"),
            "fuzzy_prompt": parsed.get("fuzzy_prompt")
        }
    
# ---- Vanna SQL Node ----
def vanna_node(state: GraphState) -> GraphState:

    if not state.get("vanna_prompt"):
        return {**state, "sql_result": pd.DataFrame([{"Error": "No vanna_prompt provided."}])}
        
    sql_query = vn_model.generate_sql(state['vanna_prompt'])
    result = vn_model.run_sql(sql_query)

    try:
        if isinstance(result, pd.DataFrame):
            parsed_result = result
        elif isinstance(result, list):
            parsed_result = pd.DataFrame(result)
        else:
            parsed_result = pd.DataFrame([{"Result": str(result)}])
    except Exception as e:
        parsed_result = pd.DataFrame([{"Error": f"Unable to parse result: {e}"}])

    print("[DEBUG] vanna output type:", type(result))
    print("[DEBUG] parsed_result type:", type(parsed_result))
    print("[DEBUG] parsed_result head:", parsed_result.head() if isinstance(parsed_result, pd.DataFrame) else parsed_result)
    
    return {**state, "sql_result": parsed_result}

# --- SerpAPI Node --- 
def serp_node(state: GraphState) -> GraphState:
    # Add insurance-specific context to the query
    query = f"{state['user_prompt']} insurance"

    search = GoogleSearch({
        "q": query,
        "api_key": os.getenv("SERPAPI_API_KEY"),
        "num": 5
    })
    results = search.get_dict()

    links = []
    if "organic_results" in results:
        for r in results["organic_results"][:5]:
            link = r.get("link")
            title = r.get("title", "Untitled")
            if link:
                links.append(f"[{title}]({link})")

    if not links:
        links = ["No insurance-related results found or API limit reached."]

    return {
        **state,
        "web_links": links
    }

# # ---- SerpAPI Search Node ----
# def serp_node(state: GraphState) -> GraphState:
#     query = state["user_prompt"]
#     search = GoogleSearch({
#         "q": query,
#         "api_key": os.getenv("SERPAPI_API_KEY"),
#         "num": 5
#     })
#     results = search.get_dict()

#     links = []
#     if "organic_results" in results:
#         for r in results["organic_results"][:5]:
#             link = r.get("link")
#             if link:
#                 links.append(link)

#     # Fallback if no results found
#     if not links:
#         links = ["No results found or API limit reached."]

#     return {
#         **state,
#         "web_links": links
#     }

# ---- Document Table Update Node ----
def document_node(state: GraphState) -> GraphState:
    doc_path = state['document_path']
    doc = Document(doc_path)

    structure_string = ""
    header = None
    header_table_map = {}

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            header = para.text.strip()
            structure_string += f"\n# {header}"
            header_table_map[header] = []
        elif header:
            header_table_map[header].append(len(header_table_map[header]))

    for idx, table in enumerate(doc.tables):
        cols = [cell.text.strip() for cell in table.rows[0].cells]
        structure_string += f"\n- Table {idx}: {len(table.rows)} rows x {len(cols)} columns, Columns: {cols}"

    prompt = f"""
            You are helping identify the correct table to update in a Word document.
            Each table has: index, rows x cols, and list of column headers.

            Document structure:
            {structure_string}

            Instruction:
            \"\"\"{state['fuzzy_prompt']}\"\"\"

            Return strictly in JSON:
            {{ "header_text": "...", "table_index_under_header": 0 }}
            """
    llm_output = call_llm(prompt)
    json_match = re.search(r'{.*}', llm_output, re.DOTALL)
    parsed = json.loads(json_match.group()) if json_match else {"header_text": list(header_table_map)[0], "table_index_under_header": 0}

    vanna_output = vn_model.run_sql(vn_model.generate_sql(state['vanna_prompt']))

    header = parsed['header_text']
    table_idx = parsed['table_index_under_header']
    matched_table_index = list(header_table_map[header])[table_idx]
    table = doc.tables[matched_table_index]

    for i in range(len(table.rows)):
        for j in range(len(table.columns)):
            if i < len(vanna_output) and j < len(vanna_output[i]):
                table.cell(i, j).text = str(vanna_output[i][j])

    updated_path = "updated_doc.docx"
    doc.save(updated_path)
    return {**state, "updated_doc_path": updated_path, "header_updated": header, "table_index_updated": matched_table_index}

# ---- LangGraph Setup ----
graph_builder = StateGraph(GraphState)
graph_builder.add_node("router", RouterNode())
graph_builder.add_node("vanna_sql", vanna_node)
graph_builder.add_node("serp_search", serp_node)
graph_builder.add_node("doc_update", document_node)

def router_logic(state: GraphState):
    if state['route'] == 'sql': return "vanna_sql"
    elif state['route'] == 'search': return "serp_search"
    elif state['route'] == 'document': return "doc_update"
    else: return END

graph_builder.set_entry_point("router")
graph_builder.add_conditional_edges("router", router_logic)
graph_builder.add_edge("vanna_sql", END)
graph_builder.add_edge("serp_search", END)
graph_builder.add_edge("doc_update", END)

agent_graph = graph_builder.compile()

# ---- Streamlit UI ----
st.title("\U0001F9E0 Agentic AI Assistant (Insurance)")

user_prompt = st.text_input("Enter your query:")
doc_file = st.file_uploader("Upload Insurance Document (.docx)", type=["docx"])

if st.button("Run Agent"):
    if doc_file:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(doc_file.read())
            doc_path = tmp.name
    else:
        doc_path = None

    state: GraphState = {
        "user_prompt": user_prompt,
        "doc_loaded": doc_path is not None,
        "document_path": doc_path,
        "vanna_prompt": None,
        "fuzzy_prompt": None,
        "route": None,
        "sql_result": None,
        "web_links": None,
        "updated_doc_path": None
    }

    with st.spinner("Running Agent..."):
        try:
            output = agent_graph.invoke(state)
        except Exception as e:
            st.error(f"Agent crashed due to error: {e}")
            import traceback
            st.code(traceback.format_exc())
            st.stop()
            
    if output.get("sql_result") is not None:
        st.subheader("SQL Query Result:")
        try:
            sql_df = output["sql_result"]
            if isinstance(sql_df, pd.DataFrame):
                st.dataframe(sql_df)
            else:
                st.write("Raw SQL output:")
                st.write(sql_df)
        except Exception as e:
            st.warning(f"Could not display table properly: {e}")
            st.write(output["sql_result"])

    if output.get("web_links"):
        st.subheader("Top Web Links:")
        for link in output["web_links"]:
            st.markdown(f"- [{link}]({link})")

    if output.get("updated_doc_path"):
        with open(output["updated_doc_path"], "rb") as f:
            st.download_button("Download Updated Document", f, file_name="updated.docx")